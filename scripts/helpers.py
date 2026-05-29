"""Public API for driving the 清华综合论文训练 Word template via python-docx.

Quick start:

    from tsinghua_thesis import helpers as h
    doc = h.open_template()
    h.set_cover_info(doc, title_cn="...", author="...", department="...",
                     major="...", advisor="...", date="二○二六年六月")
    h.set_abstract(doc, cn_text="...", cn_keywords=["k1","k2"],
                   en_text="...", en_keywords=["k1","k2"])
    h.add_chapter(doc, "引言")
    h.add_body(doc, "第一段正文。")
    h.save(doc, "thesis.docx")

After saving, open the .docx in Word/WPS and press F9 to refresh 插图清单/附表清单
TOC fields.
"""
from __future__ import annotations
import os
import re
from pathlib import Path
from typing import Literal, Optional, Sequence

from docx import Document as _DocumentCtor
from docx.document import Document
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import styles as S
from ._xml import (
    find_paragraph_by_text,
    iter_paragraphs_after,
    replace_paragraph_text,
    insert_paragraph_after,
    remove_paragraph,
    set_table_cell_text,
    append_omml,
    add_field_run,
    enable_update_fields_on_open,
    set_caption_field_cache,
)


# ---------------------------------------------------------------------------
# Template resolution
# ---------------------------------------------------------------------------

_TEMPLATE_PATH_FILE = Path(__file__).resolve().parent.parent / "template" / "TEMPLATE_PATH.txt"
_DOWNLOAD_HINT = (
    "Please obtain the official 清华综合论文训练论文模板.docx from "
    "the 清华大学教务处 download page, then configure its path by either:\n"
    "  1. passing template_path=Path(...) explicitly,\n"
    "  2. setting env var TSINGHUA_THESIS_TEMPLATE=/path/to/template.docx, or\n"
    f"  3. writing the path into {_TEMPLATE_PATH_FILE}"
)


def _resolve_template_path(template_path: Optional[Path]) -> Path:
    if template_path is not None:
        p = Path(template_path).expanduser()
        if not p.exists():
            raise FileNotFoundError(f"Template not found at explicit path: {p}")
        return p
    env = os.environ.get("TSINGHUA_THESIS_TEMPLATE")
    if env:
        p = Path(env).expanduser()
        if not p.exists():
            raise FileNotFoundError(f"TSINGHUA_THESIS_TEMPLATE points to missing file: {p}")
        return p
    if _TEMPLATE_PATH_FILE.exists():
        first_line = _TEMPLATE_PATH_FILE.read_text(encoding="utf-8").splitlines()
        # skip empty + comment lines
        for line in first_line:
            line = line.strip()
            if line and not line.startswith("#"):
                p = Path(line).expanduser()
                if not p.exists():
                    raise FileNotFoundError(
                        f"{_TEMPLATE_PATH_FILE} points to missing file: {p}"
                    )
                return p
    raise FileNotFoundError(_DOWNLOAD_HINT)


# ---------------------------------------------------------------------------
# I/O
# ---------------------------------------------------------------------------

def open_template(template_path: Optional[Path] = None) -> Document:
    """Open the official template .docx as the starting Document."""
    return _DocumentCtor(str(_resolve_template_path(template_path)))


def save(doc: Document, out_path: Path | str) -> Path:
    """Save document. Refuses to overwrite the template itself.

    Also sets `<w:updateFields val="true"/>` so Word prompts to refresh
    fields (TOC, STYLEREF, SEQ for figure/table captions, etc.) on open —
    otherwise field codes display their cached placeholder text until F9.
    """
    out = Path(out_path).expanduser().resolve()
    template = _resolve_template_path(None) if _can_resolve_template() else None
    if template is not None and out == template.resolve():
        raise ValueError(
            f"Refusing to overwrite the template at {template}. "
            "Pass a different out_path."
        )
    enable_update_fields_on_open(doc)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _can_resolve_template() -> bool:
    try:
        _resolve_template_path(None)
        return True
    except FileNotFoundError:
        return False


# ---------------------------------------------------------------------------
# Cover (Table 0 + 封面* paragraphs)
# ---------------------------------------------------------------------------

# Cover table row → label (must match what's in the template)
_COVER_TABLE_ROWS = {
    "department": 0,   # 系别
    "major":      1,   # 专业
    "author":     2,   # 姓名
    "advisor":    3,   # 指导教师
}
_COVER_TABLE_TEXT_COL = 2  # the column holding the value


def set_cover_info(
    doc: Document,
    *,
    title_cn: str,
    author: str,
    department: str,
    major: str,
    advisor: str,
    date: str,
) -> None:
    """Fill the cover page: title paragraph + 4x3 info table + date paragraph.

    Args:
        title_cn: Chinese thesis title (replaces template placeholder text in
                  the 封面论文题目 paragraph).
        author: 姓名.
        department: 系别 e.g. "自动化系".
        major: 专业 e.g. "自动化".
        advisor: e.g. "张三 教授".
        date: e.g. "二○二六年六月" (Chinese reading is conventional on cover).
    """
    # 1. Replace title paragraph (style: 封面论文题目). It's the second 封面论文题目
    #    paragraph (first one is intentionally blank in the template).
    title_paragraphs = [p for p in doc.paragraphs if p.style.name == S.COVER_TITLE]
    if not title_paragraphs:
        raise ValueError(f"Template missing any '{S.COVER_TITLE}' paragraph.")
    # Find the one with actual placeholder text; fall back to the last one
    target = next((p for p in title_paragraphs if p.text.strip()), title_paragraphs[-1])
    replace_paragraph_text(target, title_cn)

    # 2. Fill the 4x3 cover table
    if not doc.tables:
        raise ValueError("Template missing cover info table.")
    cover_table = doc.tables[0]
    field_values = {
        "department": department,
        "major": major,
        "author": author,
        "advisor": advisor,
    }
    for field, row_idx in _COVER_TABLE_ROWS.items():
        cell = cover_table.rows[row_idx].cells[_COVER_TABLE_TEXT_COL]
        set_table_cell_text(cell, field_values[field])

    # 3. Fill the date paragraph (style 封面作者信息, contains '年' in the placeholder)
    for p in doc.paragraphs:
        if p.style.name == S.COVER_AUTHOR_INFO and "年" in p.text:
            replace_paragraph_text(p, date)
            break


# ---------------------------------------------------------------------------
# Abstract / Keywords
# ---------------------------------------------------------------------------

def set_abstract(
    doc: Document,
    *,
    cn_text: str,
    cn_keywords: Sequence[str],
    en_text: str,
    en_keywords: Sequence[str],
) -> None:
    """Replace the 摘要 and Abstract sections (text + keyword line) with user content.

    `cn_text` and `en_text` may contain newlines; each line becomes a paragraph
    styled `段落` (matching the template's existing abstract paragraphs).
    """
    _replace_abstract_section(
        doc,
        heading_text="摘    要",
        new_body=cn_text,
        new_keywords_line="关键词：" + "；".join(cn_keywords),
        keyword_marker="关键词：",
    )
    _replace_abstract_section(
        doc,
        heading_text="Abstract",
        new_body=en_text,
        new_keywords_line="Keywords: " + "; ".join(en_keywords),
        keyword_marker="Keywords:",
    )


def _replace_abstract_section(
    doc: Document,
    *,
    heading_text: str,
    new_body: str,
    new_keywords_line: str,
    keyword_marker: str,
) -> None:
    heading = find_paragraph_by_text(doc, heading_text)
    if heading is None or heading.style.name != S.CHAPTER_TITLE:
        # Re-search ignoring style if the strict match failed
        heading = find_paragraph_by_text(doc, heading_text.strip())
    if heading is None:
        raise ValueError(f"Could not find abstract heading: {heading_text!r}")

    # Collect existing body paragraphs between heading and keyword line.
    body_paragraphs = []
    keyword_paragraph = None
    for p in iter_paragraphs_after(doc, heading):
        if keyword_marker in p.text:
            keyword_paragraph = p
            break
        if p.style.name == S.CHAPTER_TITLE:
            # Reached the next section without finding keywords; abort defensively
            break
        body_paragraphs.append(p)

    # Strategy: rewrite the first body paragraph to first line of new content;
    # for additional lines, insert siblings before the keyword paragraph.
    lines = [ln for ln in new_body.split("\n") if ln.strip()]
    if not lines:
        lines = [""]

    if body_paragraphs:
        # Rewrite the first body paragraph; insert remaining lines after it;
        # delete any leftover original-template body paragraphs entirely.
        replace_paragraph_text(body_paragraphs[0], lines[0])
        anchor = body_paragraphs[0]
        for line in lines[1:]:
            anchor = insert_paragraph_after(anchor, line, style=S.BODY_LEGACY)
        for extra in body_paragraphs[1:]:
            remove_paragraph(extra)

    if keyword_paragraph is not None:
        replace_paragraph_text(keyword_paragraph, new_keywords_line)


# ---------------------------------------------------------------------------
# Body structure
# ---------------------------------------------------------------------------

def add_chapter(doc: Document, title: str) -> Paragraph:
    """Append a *front-matter* chapter title (style `章标题-无级别`).

    Use for: 摘要, Abstract, 插图清单, 附表清单, 符号和缩略语说明,
    综合论文训练记录表. These are not auto-numbered by Word.

    For body chapters (引言 / 第 N 章) use `add_body_chapter` instead so that
    Word's outline numbering produces "第 1 章 引言" automatically.
    For back-matter chapters (参考文献 / 致谢 / 声明 / 在学期间研究成果) the
    template already has Title-styled headings; do not add new ones.
    """
    return doc.add_paragraph(title, style=S.CHAPTER_TITLE)


def add_body_chapter(doc: Document, title: str) -> Paragraph:
    """Append a *body* chapter title (style `Heading 1`).

    Use for 引言 and 第 N 章. Pass only the chapter name (e.g. "引言",
    "建模分析") — Word's auto-numbering adds the "第 N 章" prefix. Prefixing
    the text yourself produces duplicated numbering like "第 2 章 2. 建模...".
    """
    return doc.add_paragraph(title, style=S.HEADING_1)


def add_section(doc: Document, title: str, level: Literal[1, 2, 3, 4] = 1) -> Paragraph:
    """Append a section heading. level 1..4 → Heading 1..4.

    Pass only the section name (e.g. "聚烯烃排产问题描述"); Word auto-numbers
    Heading 2/3/4 as N.M, N.M.K, N.M.K.L within the enclosing body chapter.
    Note: `level=1` is equivalent to `add_body_chapter` and is the right
    choice when this section starts a new chapter.
    """
    style_name = {1: S.HEADING_1, 2: S.HEADING_2, 3: S.HEADING_3, 4: S.HEADING_4}[level]
    return doc.add_paragraph(title, style=style_name)


def add_body(doc: Document, text: str) -> Paragraph:
    """Append a body paragraph styled `论文正文段落`."""
    return doc.add_paragraph(text, style=S.BODY)


# ---------------------------------------------------------------------------
# Floats: Figures
# ---------------------------------------------------------------------------

def add_figure(
    doc: Document,
    image_path: Path | str,
    caption: str,
    *,
    width_cm: float = 12.0,
    label: Optional[str] = None,
) -> Paragraph:
    """Insert a figure and its caption.

    The figure occupies a `图片`-styled paragraph; the caption is a
    `Caption`-styled paragraph below it.

    Numbering:
    * `label=None` (default) → caption is auto-numbered via Word field codes,
      producing "图 N.M  caption" where N = current chapter number (via
      STYLEREF Heading 1) and M = per-chapter figure counter (via SEQ 图).
      Word displays a placeholder until F9 refreshes the fields.
    * `label="图 X-Y"` (or anything) → literal prefix, no field codes.
      Use this only for special cases (appendix figures with manual labels,
      front-matter illustrations, etc.) where auto-numbering is undesirable.
    """
    img_p = doc.add_paragraph(style=S.FIGURE_PARAGRAPH)
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap_p = doc.add_paragraph(style=S.FIG_CAPTION)
    if label:
        cap_p.add_run(f"{label}  {caption}")
    else:
        cap_p.add_run("图 ")
        add_field_run(cap_p, "STYLEREF 1 \\s", placeholder="1")
        cap_p.add_run(".")
        add_field_run(cap_p, "SEQ 图 \\* ARABIC \\s 1", placeholder="1")
        cap_p.add_run(f"  {caption}")
    return cap_p


# ---------------------------------------------------------------------------
# Floats: Three-line table
# ---------------------------------------------------------------------------

def add_three_line_table(
    doc: Document,
    header: Sequence[str],
    rows: Sequence[Sequence[str]],
    caption: str,
    *,
    label: Optional[str] = None,
) -> Table:
    """Build a 三线表-styled table with its caption above (表-题注).

    Numbering: same convention as `add_figure`.
    * `label=None` → caption auto-numbered as "表 N.M  caption" via Word
      field codes (STYLEREF Heading 1 + SEQ 表). User presses F9 to refresh.
    * `label="表 X-Y"` → literal prefix, no field codes.
    """
    cap_p = doc.add_paragraph(style=S.TABLE_CAPTION)
    if label:
        cap_p.add_run(f"{label}  {caption}")
    else:
        cap_p.add_run("表 ")
        add_field_run(cap_p, "STYLEREF 1 \\s", placeholder="1")
        cap_p.add_run(".")
        add_field_run(cap_p, "SEQ 表 \\* ARABIC \\s 1", placeholder="1")
        cap_p.add_run(f"  {caption}")

    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = S.THREE_LINE_TABLE
    # Header row
    for c, h in enumerate(header):
        table.rows[0].cells[c].text = h
    # Body rows
    for r, row in enumerate(rows, start=1):
        if len(row) != ncols:
            raise ValueError(
                f"row {r-1} has {len(row)} cells but header has {ncols} columns"
            )
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = str(val)
    return table


# ---------------------------------------------------------------------------
# Equations (pandoc-backed OMML)
# ---------------------------------------------------------------------------

def add_equation(
    doc: Document,
    latex: str,
    *,
    label: Optional[str] = None,
    on_error: Literal["raise", "text"] = "raise",
) -> Paragraph:
    """Insert a display equation as native Word OMML, styled `公式`.

    `label` (e.g. "(2-3)") is appended as trailing text in the same paragraph
    so it shows on the right; you can use Word's tab stops to right-align it.

    `\\atop` (plain-TeX stacking, common in GDP disjunction blocks) is rewritten
    to \\substack automatically so pandoc can convert it.

    on_error:
      * "raise" (default) — propagate ValueError/RuntimeError if pandoc can't
        convert the LaTeX (preserves the strict contract).
      * "text" — on failure, emit a 公式-styled paragraph containing the raw
        LaTeX as text instead of crashing. Useful when bulk-assembling many
        equations where one exotic macro shouldn't abort the whole document.

    Requires pandoc on PATH. Raises RuntimeError with install hint if absent.
    """
    from ._equation import latex_to_omml

    try:
        omaths = latex_to_omml(latex, display=True)
    except Exception:
        if on_error == "text":
            p = doc.add_paragraph(style=S.EQUATION)
            p.add_run(latex + ("\t" + label if label else ""))
            return p
        raise

    p = doc.add_paragraph(style=S.EQUATION)
    for omath in omaths:
        append_omml(p, omath)
    if label:
        # tab + label for right-alignment if user has set a right tab stop
        p.add_run("\t" + label)
    return p


def add_inline_equation(paragraph: Paragraph, latex: str) -> None:
    """Embed an inline equation inside an existing paragraph (OMML, not text)."""
    from ._equation import latex_to_omml

    for omath in latex_to_omml(latex, display=False):
        append_omml(paragraph, omath)


# --- Markdown-ish inline markup → styled runs --------------------------------

_RICH_TOKEN_RE = re.compile(
    r"\$(?P<m1>[^$\n]+?)\$"
    r"|\\\((?P<m2>[^\n]+?)\\\)"
    r"|\*\*(?P<bold>[^*\n]+?)\*\*"
    r"|`(?P<code>[^`\n]+?)`"
)
_RICH_MATH_ONLY_RE = re.compile(r"(\$[^$\n]+?\$|\\\([^\n]+?\\\))")
_RICH_MATH_PICK_RE = re.compile(r"\$([^$\n]+?)\$|\\\(([^\n]+?)\\\)")


def _render_inline_markup(paragraph: Paragraph, text: str) -> None:
    """Render `text` into `paragraph`, converting inline markup to runs/OMML:
    `$..$` / `\\(..\\)` → inline OMML; `**..**` → bold run (math inside it is
    still rendered as OMML); `` `..` `` → 行内代码 character style.
    """
    if not any(tok in text for tok in ("$", "\\(", "**", "`")):
        paragraph.add_run(text)
        return
    pos = 0
    for m in _RICH_TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("m1") is not None or m.group("m2") is not None:
            add_inline_equation(paragraph, (m.group("m1") or m.group("m2")).strip())
        elif m.group("bold") is not None:
            # bold may itself wrap inline math
            for part in _RICH_MATH_ONLY_RE.split(m.group("bold")):
                mm = _RICH_MATH_PICK_RE.fullmatch(part)
                if mm:
                    add_inline_equation(paragraph, (mm.group(1) or mm.group(2)).strip())
                elif part:
                    paragraph.add_run(part).bold = True
        else:  # code
            run = paragraph.add_run(m.group("code"))
            try:
                run.style = paragraph.part.document.styles[S.INLINE_CODE]
            except KeyError:
                pass
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_rich_body(doc: Document, text: str) -> Paragraph:
    """Like `add_body` but renders inline `$math$` / `\\(math\\)` / `**bold**`
    / `` `code` `` markup into OMML and styled runs. Plain text with no markup
    behaves exactly like `add_body`.
    """
    p = doc.add_paragraph(style=S.BODY)
    _render_inline_markup(p, text)
    return p


# ---------------------------------------------------------------------------
# Code
# ---------------------------------------------------------------------------

def add_code_block(
    doc: Document,
    code: str,
    language: Optional[str] = None,
) -> Paragraph:
    """Append code styled `行间代码`. Each input line → one paragraph.

    The `language` argument is currently recorded as a leading comment in the
    first paragraph for context; styling is otherwise uniform.
    """
    lines = code.split("\n")
    first_p = None
    if language:
        first_p = doc.add_paragraph(f"# language: {language}", style=S.CODE_BLOCK)
    for line in lines:
        p = doc.add_paragraph(line if line else " ", style=S.CODE_BLOCK)
        if first_p is None:
            first_p = p
    return first_p  # type: ignore[return-value]


def add_inline_code(paragraph: Paragraph, code: str) -> None:
    """Append a run with `行内代码` character style to an existing paragraph."""
    run = paragraph.add_run(code)
    # Apply character style
    char_style = paragraph.part.document.styles[S.INLINE_CODE]
    run.style = char_style


# ---------------------------------------------------------------------------
# References & Appendix
# ---------------------------------------------------------------------------

def add_reference(doc: Document, entry: str) -> Paragraph:
    """Append one reference entry styled `参考文献`.

    Entries should already be formatted per GB/T 7714—2015; this skill does
    not parse BibTeX (see references/known-limitations.md).

    Auto-numbering note: the `参考文献` style auto-numbers as **"[N]"**, same
    trap as `add_body_chapter` and `add_appendix_heading`. Pass only the
    citation body; do NOT prefix "[1] " yourself or Word renders "[1] [1] ...".

    ❌  add_reference(doc, "[1] 葛凌生. 产业链的多元化...")
    ✅  add_reference(doc, "葛凌生. 产业链的多元化...")

    Append references in citation order (first-cited first); Word numbers
    them in document order, so the in-text [N] citations only line up if you
    append in the same order.
    """
    return doc.add_paragraph(entry, style=S.REFERENCE)


def add_appendix_heading(
    doc: Document,
    title: str,
    level: Literal[0, 1, 2, 3] = 0,
) -> Paragraph:
    """Append `附录标题` (level=0) or `附录标题 1/2/3`.

    Auto-numbering note — same caveat as `add_body_chapter`:

    * `附录标题`   (level=0) — Word auto-numbers as **"附录 X"**.
    * `附录标题 1` (level=1) — NOT auto-numbered; pass the full text.
    * `附录标题 2` (level=2) — auto-numbered as **"X.N"**.
    * `附录标题 3` (level=3) — auto-numbered as **"X.N.M"**.

    For auto-numbered levels (0/2/3) pass ONLY the title text; do not prefix
    "附录 X " or "N. " yourself. Otherwise Word produces duplicated labels
    like "附录 A 附录 A 序列式 MILP 基准模型" or "1. 1. 初始约束".

    ❌  add_appendix_heading(doc, "附录 A 序列式 MILP 基准模型", level=0)
    ✅  add_appendix_heading(doc, "序列式 MILP 基准模型", level=0)
    """
    style_name = {
        0: S.APPENDIX_HEADING_0,
        1: S.APPENDIX_HEADING_1,
        2: S.APPENDIX_HEADING_2,
        3: S.APPENDIX_HEADING_3,
    }[level]
    return doc.add_paragraph(title, style=style_name)


def add_symbols_table(
    doc: Document,
    rows: Sequence[tuple[str, str]],
) -> Paragraph:
    """Append entries to the 符号和缩略语说明表.

    Each row is one paragraph styled `符号和缩略语说明表` with a tab between
    symbol and description (the template's existing rows use this layout).
    """
    last = None
    for sym, desc in rows:
        text = f"{sym}\t{desc}"
        last = doc.add_paragraph(text, style=S.SYMBOLS_TABLE)
    if last is None:
        raise ValueError("add_symbols_table: empty rows")
    return last


# ---------------------------------------------------------------------------
# TOC of figures / tables — placeholder helper
# ---------------------------------------------------------------------------

def insert_toc_placeholder(
    doc: Document,
    kind: Literal["figures", "tables"],
) -> Paragraph:
    """Insert a chapter-titled placeholder for 插图清单 or 附表清单.

    The actual TOC field already exists in the template; this helper is for
    cases where you've built a doc from scratch and want a labeled spot for
    the user to insert the field manually in Word (Insert > Table of Figures).
    """
    title = "插图清单" if kind == "figures" else "附表清单"
    heading = doc.add_paragraph(title, style=S.CHAPTER_TITLE)
    doc.add_paragraph(
        f"[在此插入 Word > 引用 > 插入表目录，类型选择 '{title.replace('清单','')}']",
        style=S.NORMAL,
    )
    return heading


def renumber_caption_fields(
    doc: Document,
    *,
    body_heading_style: str = "Heading 1",
    appendix_heading_style: Optional[str] = None,
) -> None:
    """Bake correct "N.M" numbers into every figure/table caption's field cache.

    Captions created by `add_figure`/`add_three_line_table` (with label=None)
    carry STYLEREF + SEQ field codes whose *displayed* value is a cached "1"
    until Word refreshes fields (F9). Until then every caption — and the
    插图清单/附表清单 that aggregate them — reads "图 1.1" / "表 1.1".

    This walks the document in order, tracks the current chapter, and rewrites
    each caption's two cached values so the document reads correctly *without*
    a refresh:

    * paragraphs styled `body_heading_style` (default "Heading 1") advance a
      numeric chapter counter (1, 2, 3, …);
    * paragraphs styled `appendix_heading_style` (e.g. "附录标题") advance a
      letter chapter label (A, B, C, …);
    * figure (`Caption`) and table (`表-题注`) counters reset at each chapter.

    The STYLEREF/SEQ field codes remain intact, so an F9 refresh in Word still
    recomputes everything (and the 插图清单/附表清单 TOC fields, which always
    need one refresh to rebuild their entries).

    Pass `appendix_heading_style="附录标题"` for an appendix-only document so
    captions are numbered "图 A.1" etc.
    """
    chap_num = 0
    appendix_idx = -1
    fig_n = 0
    tab_n = 0
    cur_label = "1"
    for p in doc.paragraphs:
        s = p.style.name if p.style else ""
        if s == body_heading_style:
            chap_num += 1
            cur_label = str(chap_num)
            fig_n = tab_n = 0
        elif appendix_heading_style is not None and s == appendix_heading_style:
            appendix_idx += 1
            cur_label = chr(ord("A") + appendix_idx)
            fig_n = tab_n = 0
        elif s == S.FIG_CAPTION:
            fig_n += 1
            set_caption_field_cache(p, cur_label, fig_n)
        elif s == S.TABLE_CAPTION:
            tab_n += 1
            set_caption_field_cache(p, cur_label, tab_n)


# ---------------------------------------------------------------------------
# Anchored insertion: put new content into the *body* of the template
# instead of appending past the back matter.
# ---------------------------------------------------------------------------
#
# The default python-docx semantics is "append" — every helper above adds
# its paragraph(s) at the end of the document. But the template already has
# back matter (参考文献 / 致谢 / 声明 / 在学期间研究成果 / 训练记录表) at the
# end, so plain appending puts your body chapters *after* the back matter.
#
# `AnchorInserter` wraps any append-style helper call so the just-added
# elements are immediately moved to right before a fixed anchor element
# (usually the 参考文献 Title paragraph). After clearing the template's
# example body chapters with `clear_example_body`, you get the correct
# Tsinghua layout:
#
#   [cover → 摘要 → ... → 符号缩略语]  (template defaults, untouched)
#   [your body chapters]                (added via AnchorInserter)
#   [参考文献 → 附录 → 致谢 → ...]      (template defaults, untouched)
#
# IMPLEMENTATION NOTE: we compare lxml elements with `==` (which checks
# underlying C-node identity), not `id()` — Python wrapper objects returned
# by `iterchildren()` are not stable across calls, so id-based set diff
# silently classifies *every* element as "new" and scrambles the document.


class AnchorInserter:
    """Call append-style helpers, then move added elements before an anchor.

    Usage::

        from scripts import helpers as h

        doc = h.open_template()
        h.clear_example_body(doc)
        anchor = h.find_chapter_anchor(doc, "参考文献", style="Title")
        ins = h.AnchorInserter(doc, anchor)

        ins(h.add_body_chapter, "引言")
        ins(h.add_body_chapter, "建模分析")
        ins(h.add_section, "聚烯烃排产问题描述", level=2)
        ins(h.add_body, "本文研究 ...")
        ins(h.add_figure, "fig.png", caption="...", label="图 2-1")
        ins(h.add_equation, r"E = mc^2", label="(2-1)")

        h.save(doc, "thesis.docx")
    """

    def __init__(self, doc: Document, anchor_element) -> None:
        self.doc = doc
        # Accept either a Paragraph wrapper or a raw lxml element.
        self.anchor = getattr(anchor_element, "_element", anchor_element)
        self.body = doc.element.body

    def __call__(self, fn, *args, **kwargs):
        # Snapshot existing children. Use == comparison (lxml element identity),
        # NOT id() — Python wrappers from iterchildren() are not stable.
        before = list(self.body.iterchildren())
        result = fn(self.doc, *args, **kwargs)
        new_elements = [c for c in self.body.iterchildren() if c not in before]
        for el in new_elements:
            self.anchor.addprevious(el)
        return result


def find_chapter_anchor(
    doc: Document,
    contains: str,
    *,
    style: Optional[str] = None,
) -> Paragraph:
    """Return the first paragraph whose text contains `contains` and (optionally) matches `style`.

    Common back-matter anchors in the Tsinghua template:

        find_chapter_anchor(doc, "参考文献", style="Title")
        find_chapter_anchor(doc, "致谢", style="Title")
        find_chapter_anchor(doc, "声明", style="Title")
        find_chapter_anchor(doc, "在学期间", style="Title")
        find_chapter_anchor(doc, "综合论文训练记录表", style="章标题-无级别")

    Raises:
        RuntimeError: no paragraph matches.
    """
    for p in doc.paragraphs:
        if contains in p.text and (style is None or (p.style and p.style.name == style)):
            return p
    crit = f"text contains {contains!r}" + (f", style={style!r}" if style else "")
    raise RuntimeError(f"find_chapter_anchor: no paragraph matched ({crit})")


def clear_example_body(doc: Document) -> None:
    """Remove the template's example body chapters (引言 + 图表示例).

    The official template ships with two sample Heading-1 chapters
    ("引    言" and "图、表及表达式示例") plus their sample subsections, body
    text, figure/table/equation examples. This helper wipes the range from
    the first body Heading 1 up to (but not including) the 参考文献 Title,
    leaving the cover/abstract/lists/symbols section before and the back
    matter after intact.

    Idempotent: re-running on an already-cleared doc is a no-op (the loop
    walks zero elements).
    """
    intro = None
    for p in doc.paragraphs:
        if p.style and p.style.name == S.HEADING_1:
            intro = p
            break
    if intro is None:
        return  # already cleared, or no body chapters to begin with
    ref = find_chapter_anchor(doc, "参考文献", style=S.TITLE)
    parent = intro._element.getparent()
    el = intro._element
    while el is not None and el is not ref._element:
        nxt = el.getnext()
        parent.remove(el)
        el = nxt
