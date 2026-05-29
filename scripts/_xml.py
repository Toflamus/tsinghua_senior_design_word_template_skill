"""Low-level OOXML helpers used internally by the public API.

Avoid using these from user code — the public API is in helpers.py.
"""
from __future__ import annotations
from typing import Iterable, Optional
from copy import deepcopy

from docx.document import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W = lambda tag: f"{{{W_NS}}}{tag}"
M = lambda tag: f"{{{M_NS}}}{tag}"


def find_paragraph_by_text(doc: Document, needle: str) -> Optional[Paragraph]:
    """Return the first paragraph whose stripped text equals or contains `needle`."""
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def iter_paragraphs_after(doc: Document, anchor: Paragraph) -> Iterable[Paragraph]:
    """Yield paragraphs strictly after `anchor`, in document order.

    Compares the underlying XML element (anchor._p) rather than Python object
    identity — `doc.paragraphs` returns fresh Paragraph wrappers each call.
    """
    anchor_el = anchor._p
    seen = False
    for p in doc.paragraphs:
        if seen:
            yield p
        if p._p is anchor_el:
            seen = True


def clear_paragraph(p: Paragraph) -> None:
    """Remove all runs from a paragraph but keep its style + paragraph properties."""
    # Remove every <w:r> child while preserving <w:pPr>
    for child in list(p._p):
        if child.tag == W("r"):
            p._p.remove(child)


def replace_paragraph_text(p: Paragraph, text: str) -> None:
    """Clear all runs and write a single run with `text`. Preserves the style."""
    clear_paragraph(p)
    p.add_run(text)


def append_omml(paragraph: Paragraph, omath_el: etree._Element) -> None:
    """Append a deep-copied <m:oMath> element inside `paragraph`'s body XML.

    Word renders OMML inside a paragraph as a sibling of <w:r>, not nested in one.
    """
    paragraph._p.append(deepcopy(omath_el))


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    """Insert a new paragraph immediately after `paragraph`. Mirrors a missing python-docx API.

    Uses python-docx's OxmlElement so the new <w:p> is a CT_P that supports
    style assignment.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph as P
    np = P(new_p, paragraph._parent)
    if style:
        np.style = paragraph.part.document.styles[style]
    if text:
        np.add_run(text)
    return np


def remove_paragraph(paragraph: Paragraph) -> None:
    """Detach a paragraph from its parent (removes it entirely from the document)."""
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def set_table_cell_text(cell, text: str, *, keep_style: bool = True) -> None:
    """Replace cell content with a single paragraph containing `text`.

    keep_style preserves the first paragraph's style and the cell's properties.
    """
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    first = cell.paragraphs[0]
    # Clear extra paragraphs
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)
    replace_paragraph_text(first, text)


def enable_update_fields_on_open(doc: Document) -> None:
    """Set `<w:updateFields val="true"/>` in settings.xml so Word auto-prompts
    to refresh all fields (TOC, STYLEREF, SEQ, etc.) on first open.

    Without this, field codes show their cached placeholder text until the
    user presses F9 — so "图 1.1" might read "图 ?.?" or "图 1.1" depending
    on the placeholder used in `add_field_run`. With this setting, Word
    displays a banner ("This document contains fields that may refer to
    other files. Do you want to update the fields?") on open; clicking Yes
    refreshes everything.
    """
    settings = doc.settings.element
    upd = settings.find(W("updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        settings.append(upd)
    upd.set(W("val"), "true")


def add_field_run(paragraph: Paragraph, instr_text: str, *, placeholder: str = "1") -> None:
    """Append a complete Word field code to `paragraph`.

    Constructs the 4-element OOXML sequence Word uses for a field:

        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> {instr_text} </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>{placeholder}</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>

    `placeholder` is the cached text shown until the user presses F9 in Word
    to refresh fields. Pick something visually plausible (e.g. "1").

    Used for the Tsinghua template's caption auto-numbering pattern:

        STYLEREF 1 \\s          → outputs current chapter number
        SEQ 图 \\* ARABIC \\s 1 → outputs per-chapter figure counter
        SEQ 表 \\* ARABIC \\s 1 → outputs per-chapter table counter
    """
    # begin
    r1 = OxmlElement("w:r")
    fld1 = OxmlElement("w:fldChar")
    fld1.set(W("fldCharType"), "begin")
    r1.append(fld1)
    paragraph._p.append(r1)
    # instrText (note: leading + trailing space so Word parses keyword cleanly)
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {instr_text.strip()} "
    r2.append(instr)
    paragraph._p.append(r2)
    # separate
    r3 = OxmlElement("w:r")
    fld3 = OxmlElement("w:fldChar")
    fld3.set(W("fldCharType"), "separate")
    r3.append(fld3)
    paragraph._p.append(r3)
    # placeholder text
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = placeholder
    r4.append(t)
    paragraph._p.append(r4)
    # end
    r5 = OxmlElement("w:r")
    fld5 = OxmlElement("w:fldChar")
    fld5.set(W("fldCharType"), "end")
    r5.append(fld5)
    paragraph._p.append(r5)


def set_caption_field_cache(paragraph: Paragraph, chapter_label: str, seq_no: int) -> None:
    """Rewrite the cached values of a caption paragraph's two field codes.

    A caption built by `add_figure`/`add_three_line_table` (with label=None)
    contains two fields: STYLEREF (chapter number) then SEQ (per-chapter
    counter). Each field caches a displayed value in the `<w:t>` run that
    immediately follows its 'separate' fldChar. This overwrites those two
    cached values with `chapter_label` and `seq_no` respectively, so the
    caption reads e.g. "图 2.1" / "表 A.3" *without* requiring an F9 refresh.

    No-op if the paragraph doesn't contain two such fields.
    """
    cached_ts = []
    expect = False
    for r in paragraph._p.findall(W("r")):
        fld = r.find(W("fldChar"))
        if fld is not None:
            expect = fld.get(W("fldCharType")) == "separate"
            continue
        if expect:
            t = r.find(W("t"))
            if t is not None:
                cached_ts.append(t)
            expect = False
    if len(cached_ts) >= 2:
        cached_ts[0].text = str(chapter_label)
        cached_ts[1].text = str(seq_no)
