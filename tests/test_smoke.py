"""End-to-end smoke test: build a tiny thesis, reopen it, assert styles.

Skips cleanly if no template is configured (so CI on forks doesn't fail).
Equation portion auto-skipped if pandoc is missing.
"""
import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from scripts import helpers as h  # noqa: E402
from scripts import styles as S   # noqa: E402


def _template_configured() -> bool:
    try:
        h._resolve_template_path(None)  # type: ignore[attr-defined]
        return True
    except FileNotFoundError:
        return False


pytestmark = pytest.mark.skipif(
    not _template_configured(),
    reason="Template path not configured; see SKILL.md > Configure template path.",
)


def test_smoke_end_to_end(tmp_path):
    """Build a doc using most helpers, save, reopen, and assert styles & content."""
    doc = h.open_template()

    h.set_cover_info(
        doc,
        title_cn="测试论文：smoke 验证",
        author="测试 学生",
        department="自动化系",
        major="自动化",
        advisor="测试 教授",
        date="二○二六年六月",
    )

    h.set_abstract(
        doc,
        cn_text="中文摘要测试段落 1。\n中文摘要测试段落 2。",
        cn_keywords=["关键词A", "关键词B"],
        en_text="English abstract paragraph 1.\nEnglish abstract paragraph 2.",
        en_keywords=["keyA", "keyB"],
    )

    h.add_symbols_table(doc, [("Σ", "求和"), ("MILP", "Mixed-Integer LP")])

    h.add_chapter(doc, "引言")
    h.add_section(doc, "1.1 背景", level=2)
    h.add_body(doc, "这是一段测试正文。")

    # Figure (valid 1x1 red PNG built via struct+zlib)
    import struct, zlib
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr = (struct.pack(">I", len(ihdr_data)) + b"IHDR" + ihdr_data
            + struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data) & 0xffffffff))
    raw = b"\x00\xff\x00\x00"
    compressed = zlib.compress(raw)
    idat = (struct.pack(">I", len(compressed)) + b"IDAT" + compressed
            + struct.pack(">I", zlib.crc32(b"IDAT" + compressed) & 0xffffffff))
    iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", zlib.crc32(b"IEND") & 0xffffffff)
    placeholder = tmp_path / "px.png"
    placeholder.write_bytes(sig + ihdr + idat + iend)
    h.add_figure(doc, placeholder, "测试图说明", label="图 1-1", width_cm=2.0)

    h.add_three_line_table(
        doc,
        header=["A", "B"],
        rows=[["1", "2"], ["3", "4"]],
        caption="测试表说明",
        label="表 1-1",
    )

    # Code
    code_p = h.add_body(doc, "示例内联代码 ")
    h.add_inline_code(code_p, "x = 1")
    h.add_code_block(doc, "print('hello')\nx = 1 + 2", language="python")

    # Reference
    h.add_reference(doc, "测试. 引用[M]. 北京: 出版社, 2026: 1-10.")

    # Appendix
    h.add_appendix_heading(doc, "附录 A  测试附录", level=0)
    h.add_body(doc, "附录正文测试。")

    # Save
    out = tmp_path / "smoke.docx"
    h.save(doc, out)
    assert out.exists() and out.stat().st_size > 0

    # Reopen and assert
    reopened = Document(str(out))
    styles_used = {p.style.name for p in reopened.paragraphs}

    # Critical styles must appear
    for required in (
        S.CHAPTER_TITLE,
        S.BODY,
        S.HEADING_2,
        S.FIG_CAPTION,
        S.TABLE_CAPTION,
        S.CODE_BLOCK,
        S.REFERENCE,
        S.APPENDIX_HEADING_0,
        S.SYMBOLS_TABLE,
        S.FIGURE_PARAGRAPH,
    ):
        assert required in styles_used, f"Missing style after save: {required}"

    # Cover table should now hold our author name
    cover_cells = [c.text for r in reopened.tables[0].rows for c in r.cells]
    assert "测试 学生" in cover_cells, f"Author not in cover cells: {cover_cells}"

    # Title paragraph should have been rewritten
    title_text = " ".join(
        p.text for p in reopened.paragraphs if p.style.name == S.COVER_TITLE
    )
    assert "smoke 验证" in title_text

    # Abstract keyword line rewritten
    full_text = "\n".join(p.text for p in reopened.paragraphs)
    assert "关键词A" in full_text and "关键词B" in full_text
    assert "keyA" in full_text and "keyB" in full_text


def test_smoke_equation_optional(tmp_path):
    """Test equation insertion if pandoc is available; skip otherwise."""
    import shutil
    if not shutil.which("pandoc"):
        pytest.skip("pandoc not installed")

    doc = h.open_template()
    h.set_cover_info(
        doc, title_cn="公式测试", author="A", department="B", major="C",
        advisor="D", date="二○二六年六月",
    )
    h.add_chapter(doc, "公式")
    h.add_equation(doc, r"E = mc^2", label="(1-1)")
    out = tmp_path / "eqn.docx"
    h.save(doc, out)
    assert out.exists()

    # Verify OMML present in the saved doc XML
    import zipfile
    with zipfile.ZipFile(out) as zf:
        with zf.open("word/document.xml") as f:
            xml_bytes = f.read()
    assert b"oMath" in xml_bytes, "OMML element not found in saved document"
